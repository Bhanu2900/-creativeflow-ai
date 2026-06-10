import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_docx(prompt, mood, genre, story_result=None, music_result=None, visual_result=None) -> str:
    doc = Document()

    # Title
    title = doc.add_heading('✨ CreativeFlow AI — Creative Universe', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    meta = doc.add_paragraph()
    meta.add_run(f"Prompt: ").bold = True
    meta.add_run(prompt)
    meta2 = doc.add_paragraph()
    meta2.add_run(f"Mood: ").bold = True
    meta2.add_run(f"{mood}   ")
    meta2.add_run(f"Genre: ").bold = True
    meta2.add_run(genre)
    doc.add_paragraph()

    # Story section
    if story_result:
        doc.add_heading(f"📖 Story: {story_result['title']}", level=1)
        doc.add_paragraph(story_result['story'])
        info = doc.add_paragraph()
        info.add_run(f"Framework: {story_result['framework']} | Words: {story_result['word_count']}").italic = True
        doc.add_paragraph()

    # Music section
    if music_result:
        doc.add_heading(f"🎵 Music: {music_result['song_title']}", level=1)
        details = doc.add_paragraph()
        details.add_run("Tempo: ").bold = True
        details.add_run(f"{music_result['tempo']}   ")
        details.add_run("Chords: ").bold = True
        details.add_run(f"{music_result['chord_progression']}")
        doc.add_paragraph()
        doc.add_heading("Lyrics", level=2)
        doc.add_paragraph(music_result['lyrics'])
        doc.add_paragraph()
        notes = doc.add_paragraph()
        notes.add_run("Production Notes: ").bold = True
        notes.add_run(music_result['production_notes'])
        doc.add_paragraph()

    # Visual section
    if visual_result:
        doc.add_heading(f"🎨 Visual: {visual_result['scene_title']}", level=1)
        doc.add_paragraph(visual_result['scene_description'])
        color = doc.add_paragraph()
        color.add_run("Color Story: ").bold = True
        color.add_run(visual_result['color_story'])
        doc.add_paragraph()
        doc.add_heading("Key Visual Elements", level=2)
        for el in visual_result['key_elements']:
            doc.add_paragraph(el, style='List Bullet')
        doc.add_paragraph()
        doc.add_heading("Image Generation Prompt", level=2)
        doc.add_paragraph(visual_result['image_prompt'])

    # Save
    output_path = "creativeflow_output.docx"
    doc.save(output_path)
    return output_path